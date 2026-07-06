import os
import tempfile
import logging
import uuid
from fastapi import FastAPI, Response
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from minio import Minio

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("artifact-generator")

app = FastAPI(title="Kito Artifact Generator")

MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "minio.monitoring.svc:9000").replace("http://", "").replace("https://", "")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "minioadmin")

def get_minio_client():
    return Minio(
        MINIO_ENDPOINT,
        access_key=MINIO_ACCESS_KEY,
        secret_key=MINIO_SECRET_KEY,
        secure=False
    )

@app.get("/healthz")
async def healthz():
    try:
        minio_client = get_minio_client()
        if not minio_client.bucket_exists("kito-generated-artifacts"):
            minio_client.make_bucket("kito-generated-artifacts")
    except Exception as e:
        logger.error(f"MinIO health check failed: {e}")
        return Response(status_code=500, content="MinIO degraded")
    return {"status": "healthy"}

@app.post("/generate")
async def generate_artifact(ast_data: dict, format: str):
    logger.info(f"Generating artifact in format: {format}")
    
    minio_client = get_minio_client()
    
    # Ensure generated bucket exists
    if not minio_client.bucket_exists("kito-generated-artifacts"):
        minio_client.make_bucket("kito-generated-artifacts")
        
    pages = ast_data.get("pages", [])
    
    # Generate filename
    file_id = str(uuid.uuid4())
    filename = f"document_{file_id}.{format.lower()}"
    
    with tempfile.TemporaryDirectory() as tmpdir:
        local_path = os.path.join(tmpdir, filename)
        
        if format.lower() == "docx":
            # DOCX generation
            doc = Document()
            doc.add_heading("Processed Document Output", 0)
            
            for page in pages:
                content = page.get("content", "")
                # Simple parser adding markdown paragraphs
                for line in content.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.strip():
                        doc.add_paragraph(line)
                doc.add_page_break()
                
            doc.save(local_path)
            
        else:
            # Default to PDF generation using ReportLab
            doc = SimpleDocTemplate(local_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                spaceAfter=6
            )
            h1_style = ParagraphStyle(
                'CustomH1',
                parent=styles['Heading1'],
                fontSize=18,
                leading=22,
                spaceAfter=12,
                keepWithNext=True
            )
            
            story = [Paragraph("Processed Document Output", styles['Title']), Spacer(1, 20)]
            
            for page in pages:
                content = page.get("content", "")
                for line in content.split("\n"):
                    if line.startswith("# "):
                        story.append(Paragraph(line[2:], h1_style))
                    elif line.startswith("## "):
                        story.append(Paragraph(line[3:], styles['Heading2']))
                    elif line.strip():
                        story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 15))
                
            doc.build(story)
            
        # Upload compiled artifact to MinIO
        minio_client.fput_object("kito-generated-artifacts", filename, local_path)
        logger.info(f"Uploaded compiled artifact to MinIO as {filename}")
        
    return {"status": "success", "object_name": filename}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5002)
